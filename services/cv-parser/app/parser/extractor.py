"""Layout-aware text extraction from uploaded résumé files (PDF / DOCX / TXT).

Résumés are frequently multi-column and rely on visual whitespace to convey
structure, so a naive "extract_text()" call from most PDF libraries collapses
columns together and scrambles reading order. We extract with x/y positioning
from pdfplumber and reconstruct line order top-to-bottom, left-to-right within
row bands, which keeps section headers and bullet lists intact for the
downstream segmenter.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

import docx
import pdfplumber

from ..config import get_settings
from .ocr import get_ocr_engine

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".jpg", ".jpeg", ".png"}
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}


class UnsupportedFileTypeError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


class FileSignatureMismatchError(ValueError):
    """Raised when a file's magic bytes don't match its claimed extension —
    e.g. an .exe renamed to resume.pdf. Defense in depth, not a full
    antivirus scan: it catches the cheap, common spoofing case."""


def _validate_signature(ext: str, data: bytes) -> None:
    if ext == ".pdf" and not data.startswith(b"%PDF-"):
        raise FileSignatureMismatchError(
            "File claims to be a PDF but doesn't start with a PDF signature."
        )
    # .docx is a ZIP (OOXML) container; a bare ZIP local-file-header signature
    # is a necessary (not sufficient) check, but catches anything that isn't
    # even zip-shaped.
    if ext == ".docx" and not data.startswith(b"PK\x03\x04"):
        raise FileSignatureMismatchError(
            "File claims to be a .docx but doesn't have a valid ZIP/OOXML signature."
        )
    if ext in (".jpg", ".jpeg") and not data.startswith(b"\xff\xd8\xff"):
        raise FileSignatureMismatchError(
            "File claims to be a JPEG but doesn't start with a JPEG signature."
        )
    if ext == ".png" and not data.startswith(b"\x89PNG\r\n\x1a\n"):
        raise FileSignatureMismatchError(
            "File claims to be a PNG but doesn't start with a PNG signature."
        )
    # .txt has no reliable magic bytes — skipped.


@dataclass
class ExtractedDocument:
    text: str
    page_count: int
    file_type: str
    warnings: list[str]


def _extract_pdf(data: bytes) -> tuple[str, int, list[str]]:
    warnings: list[str] = []
    lines: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
            if not words:
                text = page.extract_text()
                if text:
                    lines.extend(text.splitlines())
                continue
            lines.extend(_reconstruct_reading_order(words))
            lines.append("")

    has_text_layer = any(line.strip() for line in lines)
    if page_count and not has_text_layer:
        settings = get_settings()
        if settings.enable_ocr:
            ocr_text, ocr_warnings = get_ocr_engine().extract_text(
                data, max_pages=settings.ocr_max_pages
            )
            warnings.extend(ocr_warnings)
            if ocr_text.strip():
                warnings.append(
                    "No text layer found — this appears to be a scanned/image "
                    "PDF; text below was recovered via OCR and may contain "
                    "recognition errors."
                )
                return ocr_text, page_count, warnings
            warnings.append(
                "No extractable text layer found and OCR did not recover any "
                "text — this may be a scanned/image PDF."
            )
        else:
            warnings.append(
                "No extractable text layer found — this may be a scanned/image "
                "PDF that requires OCR (currently disabled: CV_PARSER_ENABLE_OCR=0)."
            )
    return "\n".join(lines), page_count, warnings


def _reconstruct_reading_order(words: list[dict]) -> list[str]:
    """Group words into visual lines (by y-position) then order left-to-right.

    This correctly interleaves two-column résumé layouts within a line band,
    rather than emitting the entire left column before the right column.
    """
    if not words:
        return []
    row_tolerance = 3.0
    rows: list[list[dict]] = []
    for word in sorted(words, key=lambda w: (w["top"], w["x0"])):
        placed = False
        for row in rows:
            if abs(row[0]["top"] - word["top"]) <= row_tolerance:
                row.append(word)
                placed = True
                break
        if not placed:
            rows.append([word])
    rows.sort(key=lambda row: row[0]["top"])
    out = []
    for row in rows:
        row_sorted = sorted(row, key=lambda w: w["x0"])
        out.append(" ".join(w["text"] for w in row_sorted))
    return out


def _extract_docx(data: bytes) -> tuple[str, int, list[str]]:
    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts), 1, []


def _extract_image(data: bytes) -> tuple[str, int, list[str]]:
    """A résumé uploaded directly as a JPG/PNG (e.g. a phone screenshot) has
    no text layer at all by definition — it's OCR or nothing, every time,
    unlike a PDF where OCR is only a fallback for the scanned-document case.
    """
    settings = get_settings()
    if not settings.enable_ocr:
        return (
            "",
            1,
            [
                "Image uploads require OCR to extract any text, and OCR is "
                "currently disabled (CV_PARSER_ENABLE_OCR=0)."
            ],
        )

    text, warnings = get_ocr_engine().extract_text_from_image(data)
    if text.strip():
        warnings.append(
            "Text recovered via OCR from an uploaded image and may contain "
            "recognition errors."
        )
    return text, 1, warnings


def _extract_txt(data: bytes) -> tuple[str, int, list[str]]:
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return data.decode(encoding), 1, []
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), 1, ["Fell back to lossy decoding."]


def extract_document(filename: str, data: bytes) -> ExtractedDocument:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    _validate_signature(ext, data)

    if ext == ".pdf":
        text, page_count, warnings = _extract_pdf(data)
    elif ext == ".docx":
        text, page_count, warnings = _extract_docx(data)
    elif ext in _IMAGE_EXTENSIONS:
        text, page_count, warnings = _extract_image(data)
    else:
        text, page_count, warnings = _extract_txt(data)

    text = _normalize_text(text)
    if not text.strip():
        raise EmptyDocumentError("No text could be extracted from the document.")

    return ExtractedDocument(
        text=text, page_count=page_count, file_type=ext.lstrip("."), warnings=warnings
    )


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("–", "-").replace("—", "-")
    text = text.replace("‘", "'").replace("’", "'")
    text = text.replace("“", '"').replace("”", '"')
    text = text.replace("\t", " ")
    lines = [line.rstrip() for line in text.split("\n")]
    normalized: list[str] = []
    blank_streak = 0
    for line in lines:
        if not line.strip():
            blank_streak += 1
            if blank_streak > 1:
                continue
        else:
            blank_streak = 0
        normalized.append(line)
    return "\n".join(normalized).strip()
