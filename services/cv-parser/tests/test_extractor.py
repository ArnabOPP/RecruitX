"""Tests for the PDF/DOCX/TXT extraction layer.

The original test suite only exercised the .txt code path (via the shared
sample_resume.txt fixture) — extract_document's PDF and DOCX branches, and
the file-signature validation, had zero coverage. These build minimal but
real PDF/DOCX files at test time so those paths are actually driven.
"""

from __future__ import annotations

import io

import docx
import pytest
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.config import get_settings
from app.parser.extractor import (
    EmptyDocumentError,
    FileSignatureMismatchError,
    UnsupportedFileTypeError,
    extract_document,
)
from app.parser.ocr import get_ocr_engine


def _build_pdf_bytes(lines: list[str]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    y = height - 72
    for line in lines:
        c.drawString(72, y, line)
        y -= 16
    c.save()
    return buf.getvalue()


def _build_docx_bytes(paragraphs: list[str]) -> bytes:
    document = docx.Document()
    for para in paragraphs:
        document.add_paragraph(para)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


RESUME_LINES = [
    "Jordan Lee",
    "jordan.lee@example.com",
    "EDUCATION",
    "B.Sc Computer Science, State University 2019 - 2023",
    "SKILLS",
    "Python, Docker, PostgreSQL",
]


def test_pdf_extraction_round_trip():
    data = _build_pdf_bytes(RESUME_LINES)
    doc = extract_document("resume.pdf", data)
    assert doc.file_type == "pdf"
    assert doc.page_count == 1
    assert "Jordan Lee" in doc.text
    assert "PostgreSQL" in doc.text


def test_docx_extraction_round_trip():
    data = _build_docx_bytes(RESUME_LINES)
    doc = extract_document("resume.docx", data)
    assert doc.file_type == "docx"
    assert "Jordan Lee" in doc.text
    assert "PostgreSQL" in doc.text


def test_docx_extracts_table_content():
    document = docx.Document()
    document.add_paragraph("EXPERIENCE")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Role"
    table.rows[0].cells[1].text = "Software Engineer"
    buf = io.BytesIO()
    document.save(buf)

    doc = extract_document("resume.docx", buf.getvalue())
    assert "Software Engineer" in doc.text


def test_txt_extraction_utf8():
    data = "Café Résumé\nSkills: C++".encode()
    doc = extract_document("resume.txt", data)
    assert "Café Résumé" in doc.text


def test_unsupported_extension_raises():
    with pytest.raises(UnsupportedFileTypeError):
        extract_document("resume.rtf", b"whatever")


def test_empty_pdf_text_layer_raises_or_warns():
    # A PDF with a valid signature but no actual pages/text should not
    # silently return an empty ParsedResume-worthy document.
    empty_pdf = _build_pdf_bytes([])
    with pytest.raises(EmptyDocumentError):
        extract_document("resume.pdf", empty_pdf)


def test_pdf_signature_mismatch_raises():
    with pytest.raises(FileSignatureMismatchError):
        extract_document("resume.pdf", b"not actually a pdf")


def test_docx_signature_mismatch_raises():
    with pytest.raises(FileSignatureMismatchError):
        extract_document("resume.docx", b"not actually a zip/docx")


def test_txt_has_no_signature_check():
    # .txt has no reliable magic bytes, so arbitrary content is accepted.
    doc = extract_document("resume.txt", b"just plain text, no signature")
    assert "just plain text" in doc.text


def _build_scanned_pdf(lines: list[str]) -> bytes:
    """Build a PDF whose "text" is actually a rasterized image with no
    underlying text layer — simulates a scanned/photographed résumé, which
    a normal text-layer extractor (pdfplumber) recovers nothing from."""
    img_width, img_height = 1600, 150 * len(lines) + 100
    image = Image.new("RGB", (img_width, img_height), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 48)
    except OSError:
        font = ImageFont.load_default(size=48)
    y = 40
    for line in lines:
        draw.text((40, y), line, fill="black", font=font)
        y += 150

    img_buf = io.BytesIO()
    image.save(img_buf, format="PNG")
    img_buf.seek(0)

    from reportlab.lib.utils import ImageReader

    pdf_buf = io.BytesIO()
    c = canvas.Canvas(pdf_buf, pagesize=(img_width, img_height))
    c.drawImage(ImageReader(img_buf), 0, 0, width=img_width, height=img_height)
    c.save()
    return pdf_buf.getvalue()


@pytest.fixture(autouse=True)
def _reset_ocr_caches():
    """get_settings() and get_ocr_engine() are process-lifetime lru_caches;
    clear them around each OCR test so env var changes (CV_PARSER_ENABLE_OCR)
    actually take effect instead of reusing a stale cached instance."""
    get_settings.cache_clear()
    get_ocr_engine.cache_clear()
    yield
    get_settings.cache_clear()
    get_ocr_engine.cache_clear()


def test_scanned_pdf_uses_ocr_fallback():
    engine = get_ocr_engine()
    engine._ensure_loaded()  # noqa: SLF001 - force the lazy load so .available is accurate
    if not engine.available:
        pytest.skip("Tesseract is not installed in this environment")
    data = _build_scanned_pdf(["SOFTWARE ENGINEER", "SKILLS PYTHON DOCKER"])
    doc = extract_document("scanned_resume.pdf", data)
    text_upper = doc.text.upper()
    assert "SOFTWARE" in text_upper or "ENGINEER" in text_upper
    assert any("OCR" in w for w in doc.warnings)


def test_scanned_pdf_ocr_disabled_raises_empty(monkeypatch):
    monkeypatch.setenv("CV_PARSER_ENABLE_OCR", "0")
    get_settings.cache_clear()
    get_ocr_engine.cache_clear()
    data = _build_scanned_pdf(["SOFTWARE ENGINEER"])
    with pytest.raises(EmptyDocumentError):
        extract_document("scanned_resume.pdf", data)


def test_ocr_engine_availability_reflects_tesseract_install():
    # Whatever the answer is on this host, it must not raise — the whole
    # point of the fail-soft wrapper is that a missing/broken Tesseract
    # install degrades gracefully instead of crashing the pipeline.
    engine = get_ocr_engine()
    assert isinstance(engine.available, bool)
